from abc import ABC, abstractmethod
from typing import List, Dict, Any, Optional
import asyncio
import aiofiles
import fitz  # PyMuPDF
import logging

# Initialize logger early
logger = logging.getLogger(__name__)

try:
    from unstructured.partition.pdf import partition_pdf
    from unstructured.partition.docx import partition_docx
    from unstructured.partition.text import partition_text
    from unstructured.partition.md import partition_md
    UNSTRUCTURED_AVAILABLE = True
except ImportError:
    UNSTRUCTURED_AVAILABLE = False
    logger.warning("Unstructured library not fully available, using fallback processing")

from docx import Document
import requests
from zenpy import Zenpy
from jira import JIRA
from ..config import settings
from ..models import DocumentChunk, DocumentMetadata, DocumentSource
import uuid
import re
import os
from datetime import datetime


class DataIngester(ABC):
    """Abstract base class for data ingesters."""
    
    @abstractmethod
    async def ingest(self, filters: Optional[Dict[str, Any]] = None) -> List[DocumentChunk]:
        """Ingest data and return document chunks."""
        pass


class PDFIngester(DataIngester):
    """PDF document ingester using PyMuPDF and unstructured."""
    
    async def ingest(self, filters: Optional[Dict[str, Any]] = None) -> List[DocumentChunk]:
        """Ingest PDF documents."""
        file_path = filters.get("file_path") if filters else None
        if not file_path:
            raise ValueError("file_path is required for PDF ingestion")
        
        try:
            logger.info(f"Starting PDF ingestion for: {file_path}")
            
            # Try unstructured first for better text extraction
            if UNSTRUCTURED_AVAILABLE:
                try:
                    elements = partition_pdf(filename=file_path)
                    text_content = "\n\n".join([str(element) for element in elements])
                    logger.info(f"Extracted {len(text_content)} characters using unstructured")
                except Exception as e:
                    logger.warning(f"Unstructured failed, falling back to PyMuPDF: {e}")
                    # Fallback to PyMuPDF
                    doc = fitz.open(file_path)
                    text_content = ""
                    for page_num in range(len(doc)):
                        page = doc.load_page(page_num)
                        text_content += page.get_text()
                    doc.close()
                    logger.info(f"Extracted {len(text_content)} characters using PyMuPDF")
            else:
                # Use PyMuPDF directly
                doc = fitz.open(file_path)
                text_content = ""
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    text_content += page.get_text()
                doc.close()
                logger.info(f"Extracted {len(text_content)} characters using PyMuPDF")
            
            if not text_content.strip():
                raise ValueError("No text content extracted from PDF")
            
            # Chunk the text
            chunks = self._chunk_text(text_content)
            logger.info(f"Created {len(chunks)} chunks from PDF")
            
            # Create document chunks
            document_chunks = []
            title = filters.get("title", os.path.basename(file_path))
            product_version = filters.get("product_version")
            tags = filters.get("tags", [])
            
            for i, chunk_text in enumerate(chunks):
                doc_id = str(uuid.uuid4())
                metadata = DocumentMetadata(
                    doc_id=doc_id,
                    title=title,
                    source_type=DocumentSource.PDF,
                    product_version=product_version,
                    tags=tags,
                    metadata={
                        "section": f"Page section {i+1}",
                        "file_path": file_path
                    }
                )
                
                chunk = DocumentChunk(
                    chunk_id=str(uuid.uuid4()),
                    doc_id=doc_id,
                    content=chunk_text,
                    metadata=metadata,
                    chunk_index=i
                )
                document_chunks.append(chunk)
            
            # Clean up temporary file
            try:
                os.unlink(file_path)
                logger.info(f"Cleaned up temporary file: {file_path}")
            except Exception as e:
                logger.warning(f"Failed to clean up temporary file {file_path}: {e}")
            
            return document_chunks
            
        except Exception as e:
            logger.error(f"Error ingesting PDF {file_path}: {e}")
            # Clean up on error
            try:
                os.unlink(file_path)
            except:
                pass
            raise
    
    def _chunk_text(self, text: str) -> List[str]:
        """Chunk text into smaller pieces."""
        # Simple sentence-based chunking
        sentences = re.split(r'[.!?]+', text)
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            
            if len(current_chunk) + len(sentence) < settings.chunk_size:
                current_chunk += sentence + ". "
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence + ". "
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks


class DocxIngester(DataIngester):
    """Word document ingester using unstructured library for better extraction."""
    
    async def ingest(self, filters: Optional[Dict[str, Any]] = None) -> List[DocumentChunk]:
        """Ingest Word documents."""
        file_path = filters.get("file_path") if filters else None
        if not file_path:
            raise ValueError("file_path is required for DOCX ingestion")
        
        try:
            logger.info(f"Starting DOCX ingestion for: {file_path}")
            
            # Try unstructured first for better text extraction
            if UNSTRUCTURED_AVAILABLE:
                try:
                    elements = partition_docx(filename=file_path)
                    text_content = "\n\n".join([str(element) for element in elements])
                    logger.info(f"Extracted {len(text_content)} characters using unstructured")
                except Exception as e:
                    logger.warning(f"Unstructured failed, falling back to python-docx: {e}")
                    # Fallback to python-docx
                    doc = Document(file_path)
                    text_content = ""
                    for paragraph in doc.paragraphs:
                        text_content += paragraph.text + "\n"
                    logger.info(f"Extracted {len(text_content)} characters using python-docx")
            else:
                # Use python-docx directly
                doc = Document(file_path)
                text_content = ""
                for paragraph in doc.paragraphs:
                    text_content += paragraph.text + "\n"
                logger.info(f"Extracted {len(text_content)} characters using python-docx")
            
            if not text_content.strip():
                raise ValueError("No text content extracted from DOCX")
            
            # Chunk the text
            chunks = self._chunk_text(text_content)
            logger.info(f"Created {len(chunks)} chunks from DOCX")
            
            # Create document chunks
            document_chunks = []
            title = filters.get("title", os.path.basename(file_path))
            product_version = filters.get("product_version")
            tags = filters.get("tags", [])
            
            for i, chunk_text in enumerate(chunks):
                doc_id = str(uuid.uuid4())
                metadata = DocumentMetadata(
                    doc_id=doc_id,
                    title=title,
                    source_type=DocumentSource.DOCX,
                    product_version=product_version,
                    tags=tags,
                    metadata={
                        "section": f"Section {i+1}",
                        "file_path": file_path
                    }
                )
                
                chunk = DocumentChunk(
                    chunk_id=str(uuid.uuid4()),
                    doc_id=doc_id,
                    content=chunk_text,
                    metadata=metadata,
                    chunk_index=i
                )
                document_chunks.append(chunk)
            
            # Clean up temporary file
            try:
                os.unlink(file_path)
                logger.info(f"Cleaned up temporary file: {file_path}")
            except Exception as e:
                logger.warning(f"Failed to clean up temporary file {file_path}: {e}")
            
            return document_chunks
            
        except Exception as e:
            logger.error(f"Error ingesting DOCX {file_path}: {e}")
            # Clean up on error
            try:
                os.unlink(file_path)
            except:
                pass
            raise
    
    def _chunk_text(self, text: str) -> List[str]:
        """Chunk text into smaller pieces."""
        # Simple paragraph-based chunking
        paragraphs = text.split('\n\n')
        chunks = []
        current_chunk = ""
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            
            if len(current_chunk) + len(paragraph) < settings.chunk_size:
                current_chunk += paragraph + "\n\n"
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = paragraph + "\n\n"
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks


class TextIngester(DataIngester):
    """Plain text file ingester."""
    
    async def ingest(self, filters: Optional[Dict[str, Any]] = None) -> List[DocumentChunk]:
        """Ingest plain text files."""
        file_path = filters.get("file_path") if filters else None
        if not file_path:
            raise ValueError("file_path is required for text ingestion")
        
        try:
            logger.info(f"Starting text file ingestion for: {file_path}")
            
            # Read the text file
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                text_content = await f.read()
            
            if not text_content.strip():
                raise ValueError("No text content found in file")
            
            # Chunk the text
            chunks = self._chunk_text(text_content)
            logger.info(f"Created {len(chunks)} chunks from text file")
            
            # Create document chunks
            document_chunks = []
            title = filters.get("title", os.path.basename(file_path))
            product_version = filters.get("product_version")
            tags = filters.get("tags", [])
            
            for i, chunk_text in enumerate(chunks):
                doc_id = str(uuid.uuid4())
                metadata = DocumentMetadata(
                    doc_id=doc_id,
                    title=title,
                    source_type=DocumentSource.TEXT,
                    product_version=product_version,
                    tags=tags,
                    metadata={
                        "section": f"Section {i+1}",
                        "file_path": file_path
                    }
                )
                
                chunk = DocumentChunk(
                    chunk_id=str(uuid.uuid4()),
                    doc_id=doc_id,
                    content=chunk_text,
                    metadata=metadata,
                    chunk_index=i
                )
                document_chunks.append(chunk)
            
            # Clean up temporary file
            try:
                os.unlink(file_path)
                logger.info(f"Cleaned up temporary file: {file_path}")
            except Exception as e:
                logger.warning(f"Failed to clean up temporary file {file_path}: {e}")
            
            return document_chunks
            
        except Exception as e:
            logger.error(f"Error ingesting text file {file_path}: {e}")
            # Clean up on error
            try:
                os.unlink(file_path)
            except:
                pass
            raise
    
    def _chunk_text(self, text: str) -> List[str]:
        """Chunk text into smaller pieces."""
        # Simple paragraph-based chunking
        paragraphs = text.split('\n\n')
        chunks = []
        current_chunk = ""
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            
            if len(current_chunk) + len(paragraph) < settings.chunk_size:
                current_chunk += paragraph + "\n\n"
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = paragraph + "\n\n"
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks


class MarkdownIngester(DataIngester):
    """Markdown file ingester."""
    
    async def ingest(self, filters: Optional[Dict[str, Any]] = None) -> List[DocumentChunk]:
        """Ingest markdown files."""
        file_path = filters.get("file_path") if filters else None
        if not file_path:
            raise ValueError("file_path is required for markdown ingestion")
        
        try:
            logger.info(f"Starting markdown ingestion for: {file_path}")
            
            # Try unstructured first for better markdown processing
            if UNSTRUCTURED_AVAILABLE:
                try:
                    elements = partition_md(filename=file_path)
                    text_content = "\n\n".join([str(element) for element in elements])
                    logger.info(f"Extracted {len(text_content)} characters using unstructured")
                except Exception as e:
                    logger.warning(f"Unstructured failed, falling back to raw text: {e}")
                    # Fallback to raw text reading
                    async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                        text_content = await f.read()
                    logger.info(f"Extracted {len(text_content)} characters as raw text")
            else:
                # Read as raw text directly
                async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                    text_content = await f.read()
                logger.info(f"Extracted {len(text_content)} characters as raw text")
            
            if not text_content.strip():
                raise ValueError("No text content extracted from markdown")
            
            # Chunk the text
            chunks = self._chunk_text(text_content)
            logger.info(f"Created {len(chunks)} chunks from markdown")
            
            # Create document chunks
            document_chunks = []
            title = filters.get("title", os.path.basename(file_path))
            product_version = filters.get("product_version")
            tags = filters.get("tags", [])
            
            for i, chunk_text in enumerate(chunks):
                doc_id = str(uuid.uuid4())
                metadata = DocumentMetadata(
                    doc_id=doc_id,
                    title=title,
                    source_type=DocumentSource.MARKDOWN,
                    product_version=product_version,
                    tags=tags,
                    metadata={
                        "section": f"Section {i+1}",
                        "file_path": file_path
                    }
                )
                
                chunk = DocumentChunk(
                    chunk_id=str(uuid.uuid4()),
                    doc_id=doc_id,
                    content=chunk_text,
                    metadata=metadata,
                    chunk_index=i
                )
                document_chunks.append(chunk)
            
            # Clean up temporary file
            try:
                os.unlink(file_path)
                logger.info(f"Cleaned up temporary file: {file_path}")
            except Exception as e:
                logger.warning(f"Failed to clean up temporary file {file_path}: {e}")
            
            return document_chunks
            
        except Exception as e:
            logger.error(f"Error ingesting markdown {file_path}: {e}")
            # Clean up on error
            try:
                os.unlink(file_path)
            except:
                pass
            raise
    
    def _chunk_text(self, text: str) -> List[str]:
        """Chunk text into smaller pieces for markdown."""
        # Try to chunk by headers first, then fall back to paragraphs
        header_chunks = re.split(r'\n(?=#{1,6}\s)', text)
        
        if len(header_chunks) > 1:
            # We have headers, use them for chunking
            chunks = []
            for header_chunk in header_chunks:
                header_chunk = header_chunk.strip()
                if not header_chunk:
                    continue
                
                if len(header_chunk) <= settings.chunk_size:
                    chunks.append(header_chunk)
                else:
                    # Split large header sections by paragraphs
                    para_chunks = self._chunk_by_paragraphs(header_chunk)
                    chunks.extend(para_chunks)
            return chunks
        else:
            # No headers found, use paragraph-based chunking
            return self._chunk_by_paragraphs(text)
    
    def _chunk_by_paragraphs(self, text: str) -> List[str]:
        """Chunk text by paragraphs."""
        paragraphs = text.split('\n\n')
        chunks = []
        current_chunk = ""
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            
            if len(current_chunk) + len(paragraph) < settings.chunk_size:
                current_chunk += paragraph + "\n\n"
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = paragraph + "\n\n"
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks


class ZendeskIngester(DataIngester):
    """Zendesk ticket ingester."""
    
    def __init__(self):
        self.zenpy_client = None
    
    def _get_zenpy_client(self):
        """Lazy initialization of Zenpy client."""
        if self.zenpy_client is None:
            if not all([settings.zendesk_subdomain, settings.zendesk_email, settings.zendesk_token]):
                raise ValueError("Zendesk credentials are required")
            
            self.zenpy_client = Zenpy(
                subdomain=settings.zendesk_subdomain,
                email=settings.zendesk_email,
                token=settings.zendesk_token
            )
        return self.zenpy_client
    
    async def ingest(self, filters: Optional[Dict[str, Any]] = None) -> List[DocumentChunk]:
        """Ingest Zendesk tickets."""
        try:
            chunks = []
            
            # Get tickets (limit for demo purposes)
            zenpy_client = self._get_zenpy_client()
            tickets = list(zenpy_client.tickets(limit=filters.get("limit", 100)))
            
            for ticket in tickets:
                # Create metadata
                metadata = DocumentMetadata(
                    doc_id=f"zendesk_ticket_{ticket.id}",
                    title=f"Zendesk Ticket #{ticket.id}: {ticket.subject}",
                    source_type=DocumentSource.ZENDESK,
                    product_version=filters.get("product_version"),
                    tags=ticket.tags if hasattr(ticket, 'tags') else [],
                    metadata={
                        "ticket_id": str(ticket.id),
                        "status": ticket.status,
                        "priority": ticket.priority,
                        "requester_id": ticket.requester_id,
                        "assignee_id": ticket.assignee_id,
                        "created_at": ticket.created_at.isoformat() if ticket.created_at else None,
                        "updated_at": ticket.updated_at.isoformat() if ticket.updated_at else None,
                    }
                )
                
                # Combine ticket content
                content_parts = [
                    f"Subject: {ticket.subject}",
                    f"Description: {ticket.description or 'No description'}"
                ]
                
                # Add comments
                comments = list(self.zenpy_client.tickets.comments(ticket.id))
                for comment in comments[:5]:  # Limit comments for chunking
                    if comment.body:
                        content_parts.append(f"Comment: {comment.body}")
                
                full_content = "\n\n".join(content_parts)
                
                # Chunk the content
                text_chunks = self._chunk_text(full_content)
                
                # Create DocumentChunk objects
                for i, chunk_text in enumerate(text_chunks):
                    chunk = DocumentChunk(
                        chunk_id=str(uuid.uuid4()),
                        doc_id=metadata.doc_id,
                        content=chunk_text.strip(),
                        metadata=metadata,
                        chunk_index=i
                    )
                    chunks.append(chunk)
            
            logger.info(f"Ingested {len(tickets)} Zendesk tickets, created {len(chunks)} chunks")
            return chunks
            
        except Exception as e:
            logger.error(f"Error ingesting Zendesk tickets: {e}")
            raise
    
    def _chunk_text(self, text: str) -> List[str]:
        """Chunk text into smaller pieces."""
        if len(text) <= settings.chunk_size:
            return [text]
        
        # Split by sections (Subject, Description, Comments)
        sections = text.split('\n\n')
        chunks = []
        current_chunk = ""
        
        for section in sections:
            if len(current_chunk) + len(section) < settings.chunk_size:
                current_chunk += section + "\n\n"
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = section + "\n\n"
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks


class JiraIngester(DataIngester):
    """Jira ticket ingester."""
    
    def __init__(self):
        self.jira_client = None
    
    def _get_jira_client(self):
        """Lazy initialization of Jira client."""
        if self.jira_client is None:
            if not all([settings.jira_server, settings.jira_email, settings.jira_api_token]):
                raise ValueError("Jira credentials are required")
            
            self.jira_client = JIRA(
                server=settings.jira_server,
                basic_auth=(settings.jira_email, settings.jira_api_token)
            )
        return self.jira_client
    
    async def ingest(self, filters: Optional[Dict[str, Any]] = None) -> List[DocumentChunk]:
        """Ingest Jira tickets."""
        try:
            chunks = []
            
            # Build JQL query
            jql = filters.get("jql", "project is not EMPTY")
            max_results = filters.get("limit", 100)
            
            # Search issues
            jira_client = self._get_jira_client()
            issues = jira_client.search_issues(jql, maxResults=max_results, expand='comments')
            
            for issue in issues:
                # Create metadata
                metadata = DocumentMetadata(
                    doc_id=f"jira_issue_{issue.key}",
                    title=f"Jira Issue {issue.key}: {issue.fields.summary}",
                    source_type=DocumentSource.JIRA,
                    product_version=filters.get("product_version"),
                    tags=[issue.fields.issuetype.name] if hasattr(issue.fields, 'issuetype') else [],
                    metadata={
                        "ticket_id": issue.key,
                        "issue_type": issue.fields.issuetype.name if hasattr(issue.fields, 'issuetype') else None,
                        "status": issue.fields.status.name if hasattr(issue.fields, 'status') else None,
                        "priority": issue.fields.priority.name if hasattr(issue.fields, 'priority') else None,
                        "assignee": issue.fields.assignee.displayName if hasattr(issue.fields, 'assignee') and issue.fields.assignee else None,
                        "reporter": issue.fields.reporter.displayName if hasattr(issue.fields, 'reporter') else None,
                        "created_at": issue.fields.created,
                        "updated_at": issue.fields.updated,
                    }
                )
                
                # Combine issue content
                content_parts = [
                    f"Summary: {issue.fields.summary}",
                    f"Description: {issue.fields.description or 'No description'}"
                ]
                
                # Add comments
                if hasattr(issue.fields, 'comment') and issue.fields.comment.comments:
                    for comment in issue.fields.comment.comments[:5]:  # Limit comments
                        content_parts.append(f"Comment: {comment.body}")
                
                full_content = "\n\n".join(content_parts)
                
                # Chunk the content
                text_chunks = self._chunk_text(full_content)
                
                # Create DocumentChunk objects
                for i, chunk_text in enumerate(text_chunks):
                    chunk = DocumentChunk(
                        chunk_id=str(uuid.uuid4()),
                        doc_id=metadata.doc_id,
                        content=chunk_text.strip(),
                        metadata=metadata,
                        chunk_index=i
                    )
                    chunks.append(chunk)
            
            logger.info(f"Ingested {len(issues)} Jira issues, created {len(chunks)} chunks")
            return chunks
            
        except Exception as e:
            logger.error(f"Error ingesting Jira issues: {e}")
            raise
    
    def _chunk_text(self, text: str) -> List[str]:
        """Chunk text into smaller pieces."""
        if len(text) <= settings.chunk_size:
            return [text]
        
        # Split by sections (Summary, Description, Comments)
        sections = text.split('\n\n')
        chunks = []
        current_chunk = ""
        
        for section in sections:
            if len(current_chunk) + len(section) < settings.chunk_size:
                current_chunk += section + "\n\n"
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = section + "\n\n"
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks


class IngestionService:
    """Main ingestion service that coordinates different ingesters."""
    
    def __init__(self):
        self.ingesters = {
            DocumentSource.PDF: PDFIngester(),
            DocumentSource.DOCX: DocxIngester(),
            DocumentSource.TEXT: TextIngester(),
            DocumentSource.MARKDOWN: MarkdownIngester(),
        }
        
        # Initialize optional ingesters if credentials are available
        try:
            self.ingesters[DocumentSource.ZENDESK] = ZendeskIngester()
        except ValueError:
            logger.warning("Zendesk ingester not initialized - missing credentials")
        
        try:
            self.ingesters[DocumentSource.JIRA] = JiraIngester()
        except ValueError:
            logger.warning("Jira ingester not initialized - missing credentials")
    
    async def ingest_documents(
        self, 
        source_type: DocumentSource, 
        filters: Optional[Dict[str, Any]] = None
    ) -> List[DocumentChunk]:
        """Ingest documents from a specific source."""
        if source_type not in self.ingesters:
            raise ValueError(f"Ingester for {source_type} not available")
        
        ingester = self.ingesters[source_type]
        return await ingester.ingest(filters)
    
    async def ingest_multiple_sources(
        self, 
        sources: List[Dict[str, Any]]
    ) -> List[DocumentChunk]:
        """Ingest from multiple sources."""
        all_chunks = []
        
        for source_config in sources:
            source_type = DocumentSource(source_config["source_type"])
            filters = source_config.get("filters", {})
            
            try:
                chunks = await self.ingest_documents(source_type, filters)
                all_chunks.extend(chunks)
            except Exception as e:
                logger.error(f"Error ingesting from {source_type}: {e}")
                # Continue with other sources
        
        return all_chunks


# Global ingestion service instance
ingestion_service = IngestionService()
